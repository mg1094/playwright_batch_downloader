#!/usr/bin/env python3
"""
文档校验模块

使用OpenAI多模态API对下载的文档进行内容和格式校验。
支持的校验项目：
1. 两表格内容、样式格式是否一致
2. 材料名称和空白表格主旨是否相符  
3. 材料名称和示例样表主旨是否相符
4. 空白表格无示例
5. 示例样表包含填写示例
6. 示例样表信息是否打码
"""

import os
import base64
from io import BytesIO
from typing import Dict, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from PIL import Image
from openai import AsyncOpenAI
from PIL import ImageDraw, ImageFont
# 这个是 python-docx 包引入的
from docx import Document


@dataclass
class ValidationResult:
    """校验结果数据类"""
    forms_consistent: Optional[bool] = None  # 两表格内容、样式格式是否一致
    blank_form_matches: Optional[bool] = None  # 材料名称和空白表格主旨是否相符
    sample_form_matches: Optional[bool] = None  # 材料名称和示例样表主旨是否相符
    blank_form_empty: Optional[bool] = None  # 空白表格无示例
    sample_form_filled: Optional[bool] = None  # 示例样表包含填写示例
    sample_info_masked: Optional[bool] = None  # 示例样表信息是否打码
    
    # 详细说明
    forms_consistent_reason: str = ""
    blank_form_matches_reason: str = ""
    sample_form_matches_reason: str = ""
    blank_form_empty_reason: str = ""
    sample_form_filled_reason: str = ""
    sample_info_masked_reason: str = ""


class DocumentValidator:
    """文档校验器"""
    
    def __init__(self, openai_api_key: str, openai_base_url: str = None):
        """
        初始化文档校验器
        
        Args:
            openai_api_key: OpenAI API密钥
            openai_base_url: OpenAI API基础URL（可选，用于使用代理或第三方服务）
        """
        try:
            # 构建客户端初始化参数
            client_kwargs = {
                "api_key": openai_api_key,
            }
            
            # 只有在提供了base_url时才添加
            if openai_base_url:
                client_kwargs["base_url"] = openai_base_url
            
            # 为了兼容OpenRouter等第三方服务，添加超时配置
            client_kwargs["timeout"] = 60.0
            
            # 初始化AsyncOpenAI客户端
            self.client = AsyncOpenAI(**client_kwargs)
            
            # 测试连接可用性
            print(f"✅ OpenAI客户端初始化成功，API端点: {openai_base_url or 'https://api.openai.com'}")
            
        except Exception as e:
            print(f"❌ OpenAI客户端初始化失败: {e}")
            print(f"   API密钥: {'已设置' if openai_api_key else '未设置'}")
            print(f"   API端点: {openai_base_url or 'https://api.openai.com'}")
            # 不抛出异常，而是设置client为None，让系统继续运行但跳过校验
            self.client = None
    
    async def close(self):
        """
        关闭OpenAI客户端，释放资源
        """
        try:
            if hasattr(self, 'client') and self.client:
                await self.client.close()
        except Exception as e:
            print(f"⚠️ 关闭OpenAI客户端时出错: {e}")
        
    def convert_document_to_image(self, file_path: str, output_dir: str = "temp_images") -> str:
        """
        将文档转换为高清图片
        
        Args:
            file_path: 文档文件路径
            output_dir: 图片输出目录
            
        Returns:
            str: 生成的图片路径
            
        Raises:
            ValueError: 不支持的文件格式
            Exception: 文档转换失败
        """
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        file_ext = Path(file_path).suffix.lower()
        file_stem = Path(file_path).stem
        output_path = os.path.join(output_dir, f"{file_stem}.png")
        
        try:
            if file_ext in ['.pdf']:
                # 处理PDF文件
                pdf_doc = fitz.open(file_path)
                page = pdf_doc[0]  # 取第一页
                
                # 设置高分辨率渲染
                mat = fitz.Matrix(2.0, 2.0)  # 放大2倍提高清晰度
                pix = page.get_pixmap(matrix=mat)
                
                # 转换为PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                
                pdf_doc.close()
                
            elif file_ext in ['.doc', '.docx']:
                # 处理Word文档 - 先转换为PDF再转换为图片
                # 使用python-docx无法获得视觉效果，所以我们使用LibreOffice进行转换
                
                # 首先尝试转换为PDF
                temp_pdf_path = os.path.join(output_dir, f"{file_stem}_temp.pdf")
                
                # 使用LibreOffice命令行工具转换（需要系统安装LibreOffice）
                import subprocess
                try:
                    subprocess.run([
                        'soffice', '--headless', '--convert-to', 'pdf', 
                        '--outdir', output_dir, file_path
                    ], check=True, capture_output=True)
                    
                    # 转换生成的PDF为图片
                    generated_pdf = os.path.join(output_dir, f"{file_stem}.pdf")
                    if os.path.exists(generated_pdf):
                        pdf_doc = fitz.open(generated_pdf)
                        page = pdf_doc[0]
                        mat = fitz.Matrix(2.0, 2.0)
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        img = Image.open(BytesIO(img_data))
                        pdf_doc.close()
                        # 清理临时PDF文件
                        os.remove(generated_pdf)
                    else:
                        raise Exception("LibreOffice转换PDF失败")
                        
                except (subprocess.CalledProcessError, FileNotFoundError):
                    # LibreOffice不可用，尝试直接使用python-docx提取内容并生成简单图片
                    print("⚠️ LibreOffice不可用，使用备用方案生成文档预览")
                    img = self._create_text_image_from_docx(file_path)
                    
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            # 保存图片
            img.save(output_path, "PNG", quality=95)
            
            print(f"✅ 文档转换完成: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ 文档转换失败: {e}")
            raise
    
    def _create_text_image_from_docx(self, docx_path: str) -> Image.Image:
        """
        从DOCX文件创建简单的文本图片（备用方案）
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            PIL.Image: 生成的图片
        """
        try:
            
            
            # 读取DOCX内容
            doc = Document(docx_path)
            text_lines = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_lines.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_lines.append(" | ".join(row_text))
            
            # 创建简单的文本图片
            
            
            # 图片设置
            width, height = 800, max(600, len(text_lines) * 30 + 100)
            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)
            
            # 尝试使用系统字体，如果失败则使用默认字体
            try:
                font = ImageFont.truetype("Arial.ttc", 16)
            except:
                font = ImageFont.load_default()
            
            # 绘制文本
            y_offset = 20
            for line in text_lines[:30]:  # 最多显示30行
                # 文本换行处理
                if len(line) > 60:
                    line = line[:60] + "..."
                draw.text((20, y_offset), line, fill='black', font=font)
                y_offset += 25
            
            return img
            
        except Exception as e:
            print(f"⚠️ 创建文本图片失败: {e}")
            # 返回一个空白图片
            img = Image.new('RGB', (800, 600), color='white')
            draw = ImageDraw.Draw(img)
            draw.text((50, 50), f"无法读取文档内容: {os.path.basename(docx_path)}", fill='black')
            return img
    
    def _encode_image_to_base64(self, image_path: str) -> str:
        """
        将图片编码为base64字符串
        
        Args:
            image_path: 图片路径
            
        Returns:
            str: base64编码的图片数据
        """
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    async def validate_documents(
        self, 
        material_name: str, 
        blank_form_path: str = None, 
        sample_form_path: str = None
    ) -> ValidationResult:
        """
        校验文档内容和格式
        
        Args:
            material_name: 材料名称
            blank_form_path: 空白表格文件路径（可选）
            sample_form_path: 示例样表文件路径（可选）
            
        Returns:
            ValidationResult: 校验结果
        """
        result = ValidationResult()
        
        # 检查OpenAI客户端是否可用
        if not self.client:
            error_msg = "OpenAI客户端不可用，无法进行文档校验"
            result.forms_consistent_reason = error_msg
            result.blank_form_matches_reason = error_msg
            result.sample_form_matches_reason = error_msg
            result.blank_form_empty_reason = error_msg
            result.sample_form_filled_reason = error_msg
            result.sample_info_masked_reason = error_msg
            return result
        
        try:
            # 准备图片
            images = {}
            if blank_form_path and os.path.exists(blank_form_path):
                blank_img_path = self.convert_document_to_image(blank_form_path)
                images['blank'] = self._encode_image_to_base64(blank_img_path)
                
            if sample_form_path and os.path.exists(sample_form_path):
                sample_img_path = self.convert_document_to_image(sample_form_path)
                images['sample'] = self._encode_image_to_base64(sample_img_path)
            
            # 根据可用的文档进行不同的校验
            if 'blank' in images and 'sample' in images:
                # 两个文档都存在，进行完整校验
                await self._validate_both_documents(result, material_name, images['blank'], images['sample'])
                
            elif 'blank' in images:
                # 只有空白表格
                await self._validate_blank_document(result, material_name, images['blank'])
                
            elif 'sample' in images:
                # 只有示例样表
                await self._validate_sample_document(result, material_name, images['sample'])
            
            # 清理临时图片文件
            # if 'blank' in images:
            #     temp_blank_path = blank_img_path
            #     if os.path.exists(temp_blank_path):
            #         os.remove(temp_blank_path)
                    
            # if 'sample' in images:
            #     temp_sample_path = sample_img_path
            #     if os.path.exists(temp_sample_path):
            #         os.remove(temp_sample_path)
            
        except Exception as e:
            print(f"❌ 文档校验出错: {e}")
            # 所有校验结果都设为None，并记录错误信息
            result.forms_consistent_reason = f"校验出错: {e}"
            result.blank_form_matches_reason = f"校验出错: {e}"
            result.sample_form_matches_reason = f"校验出错: {e}"
            result.blank_form_empty_reason = f"校验出错: {e}"
            result.sample_form_filled_reason = f"校验出错: {e}"
            result.sample_info_masked_reason = f"校验出错: {e}"
        
        return result
    
    async def _validate_both_documents(
        self, 
        result: ValidationResult, 
        material_name: str, 
        blank_image_b64: str, 
        sample_image_b64: str
    ):
        """
        校验两个文档（空白表格和示例样表）
        """
        prompt = f"""
你是一位专业的政务材料审查员。请仔细分析下面两张文档图片，进行全面校验。

# 输入信息
- 材料名称: "{material_name}"
- 图片A: 空白表格
- 图片B: 示例样表

# 校验任务
请对以下6个方面进行校验，并给出明确的true/false判断：

1. `consistency`: 空白表格和示例样表在整体布局、结构、表格项目和样式上是否基本一致？它们看起来应该是同一个模板的两种状态（一个未填写，一个已填写）。
2. `blank_form_matches`: 请宽松判断：空白表格的主旨是否与材料名称“{material_name}”的核心主题紧密相关？例如，如果材料名称是“体检合格证明”，表格标题是“教师资格申请人员体检表”也应视为相符，因为它们都围绕“体检”这一核心主题。（判断标准：主题相关即可，无需文字完全匹配）
3. `sample_form_matches`: 请宽松判断：示例样表的主旨是否与材料名称“{material_name}”的核心主题紧密相关？例如，如果材料名称是“体检合格证明”，表格标题是“教师资格申请人员体检表”也应视为相符，因为它们都围绕“体检”这一核心主题。（判断标准：主题相关即可，无需文字完全匹配）
4. `blank_form_has_no_samples`: 空白表格中是否不包含任何个人信息填写示例？表格应该是干净的、待填写的状态。
5. `sample_form_has_samples`: 示例样表中是否清晰地包含了填写示例？
6. `sample_info_masked`: 示例样表中的个人信息（姓名、电话、地址等）是否已经打码处理？（如"张xx"、"139xxxx"等）

# 输出格式 (严格JSON格式)
{{
  "forms_consistent": true,
  "forms_consistent_reason": "详细说明判断理由",
  "blank_form_matches": true,
  "blank_form_matches_reason": "详细说明判断理由",
  "sample_form_matches": true,
  "sample_form_matches_reason": "详细说明判断理由",
  "blank_form_empty": true,
  "blank_form_empty_reason": "详细说明判断理由",
  "sample_form_filled": true,
  "sample_form_filled_reason": "详细说明判断理由",
  "sample_info_masked": true,
  "sample_info_masked_reason": "详细说明判断理由"
}}"""
        
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{blank_image_b64}",
                            "detail": "high"
                        }
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{sample_image_b64}",
                            "detail": "high"
                        }
                    }
                ]
            }
        ]
        
        try:
            response = await self.client.chat.completions.create(
                model="qwen-vl-max-latest",  # 使用支持视觉的模型
                messages=messages,
                max_tokens=1000,
                temperature=0.1
            )
            
            # 解析响应
            content = response.choices[0].message.content
            print("ai调用成功，开始解析响应")
            await self._parse_validation_response(result, content)
            
        except Exception as e:
            print(f"❌ OpenAI API调用失败: {e}")
            error_msg = f"API调用失败: {e}"
            result.forms_consistent_reason = error_msg
            result.blank_form_matches_reason = error_msg
            result.sample_form_matches_reason = error_msg
            result.blank_form_empty_reason = error_msg
            result.sample_form_filled_reason = error_msg
            result.sample_info_masked_reason = error_msg
    
    async def _validate_blank_document(
        self, 
        result: ValidationResult, 
        material_name: str, 
        blank_image_b64: str
    ):
        """
        校验单个空白表格文档
        """
        prompt = f"""
        你是一位专业的政务材料审查员。请仔细分析下面的文档图片，进行全面校验。
        
        # 输入信息
        - 材料名称: "{material_name}"
        - 文档类型: 空白表格
        
        # 校验图片
        <image>
        
        # 校验要求
        请对以下3个方面进行校验，并给出明确的true/false判断：
        
        1. `blank_form_matches`: 请宽松判断：空白表格的主旨是否与材料名称“{material_name}”的核心主题紧密相关？例如，如果材料名称是“体检合格证明”，表格标题是“教师资格申请人员体检表”也应视为相符，因为它们都围绕“体检”这一核心主题。（判断标准：主题相关即可，无需文字完全匹配）
        2. `blank_form_has_no_samples`: 空白表格中是否不包含任何个人信息填写示例？表格应该是干净的、待填写的状态。
        
        # 输出格式
{{
  "blank_form_matches": true,
  "blank_form_matches_reason": "详细说明判断理由",
  "blank_form_empty": true,
  "blank_form_empty_reason": "详细说明判断理由"
}}"""
        
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{blank_image_b64}",
                            "detail": "high"
                        }
                    }
                ]
            }
        ]
        
        try:
            response = await self.client.chat.completions.create(
                model="qwen-vl-max-latest",
                messages=messages,
                max_tokens=500,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            await self._parse_validation_response(result, content)
            
        except Exception as e:
            print(f"❌ OpenAI API调用失败: {e}")
            error_msg = f"API调用失败: {e}"
            result.blank_form_matches_reason = error_msg
            result.blank_form_empty_reason = error_msg
    
    async def _validate_sample_document(
        self, 
        result: ValidationResult, 
        material_name: str, 
        sample_image_b64: str
    ):
        """
        校验单个示例样表文档
        """
        prompt = f"""
        你是一位专业的政务材料审查员。请仔细分析下面的文档图片，进行全面校验。
        
        # 输入信息
        - 材料名称: "{material_name}"
        - 文档类型: 示例样表
        
        # 校验图片
        <image>
        
        # 校验要求
        请对以下3个方面进行校验，并给出明确的true/false判断：
        
        1. `sample_form_matches`: 请宽松判断：示例样表的主旨是否与材料名称“{material_name}”的核心主题紧密相关？例如，如果材料名称是“体检合格证明”，表格标题是“教师资格申请人员体检表”也应视为相符，因为它们都围绕“体检”这一核心主题。（判断标准：主题相关即可，无需文字完全匹配）
        2. `sample_form_has_samples`: 示例样表中是否清晰地包含了填写示例？
        3. `sample_info_masked`: 示例样表中的个人信息（姓名、电话、地址等）是否已经打码处理？（如"张xx"、"139xxxx"等）
        
        # 输出格式
{{
  "sample_form_matches": true,
  "sample_form_matches_reason": "详细说明判断理由",
  "sample_form_filled": true,
  "sample_form_filled_reason": "详细说明判断理由",
  "sample_info_masked": true,
  "sample_info_masked_reason": "详细说明判断理由"
}}"""
        
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{sample_image_b64}",
                            "detail": "high"
                        }
                    }
                ]
            }
        ]
        
        try:
            response = await self.client.chat.completions.create(
                model="qwen-vl-max-latest",
                messages=messages,
                max_tokens=500,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            await self._parse_validation_response(result, content)
            
        except Exception as e:
            print(f"❌ OpenAI API调用失败: {e}")
            error_msg = f"API调用失败: {e}"
            result.sample_form_matches_reason = error_msg
            result.sample_form_filled_reason = error_msg
            result.sample_info_masked_reason = error_msg
    
    async def _parse_validation_response(self, result: ValidationResult, content: str):
        """
        解析OpenAI API响应并更新ValidationResult
        """
        try:
            import json
            
            # 尝试提取JSON部分
            if "```json" in content:
                json_start = content.find("```json") + 7
                json_end = content.find("```", json_start)
                json_content = content[json_start:json_end].strip()
            elif "{" in content and "}" in content:
                json_start = content.find("{")
                json_end = content.rfind("}") + 1
                json_content = content[json_start:json_end]
            else:
                json_content = content
            
            # 解析JSON
            data = json.loads(json_content)
            
            # 更新结果
            if "forms_consistent" in data:
                result.forms_consistent = data["forms_consistent"]
                result.forms_consistent_reason = data.get("forms_consistent_reason", "")
                
            if "blank_form_matches" in data:
                result.blank_form_matches = data["blank_form_matches"]
                result.blank_form_matches_reason = data.get("blank_form_matches_reason", "")
                
            if "sample_form_matches" in data:
                result.sample_form_matches = data["sample_form_matches"]
                result.sample_form_matches_reason = data.get("sample_form_matches_reason", "")
                
            if "blank_form_empty" in data:
                result.blank_form_empty = data["blank_form_empty"]
                result.blank_form_empty_reason = data.get("blank_form_empty_reason", "")
                
            if "sample_form_filled" in data:
                result.sample_form_filled = data["sample_form_filled"]
                result.sample_form_filled_reason = data.get("sample_form_filled_reason", "")
                
            if "sample_info_masked" in data:
                result.sample_info_masked = data["sample_info_masked"]
                result.sample_info_masked_reason = data.get("sample_info_masked_reason", "")
            
        except Exception as e:
            print(f"❌ 解析API响应失败: {e}")
            print(f"原始响应: {content}")
            error_msg = f"解析响应失败: {e}"
            if result.forms_consistent_reason == "":
                result.forms_consistent_reason = error_msg
            if result.blank_form_matches_reason == "":
                result.blank_form_matches_reason = error_msg
            if result.sample_form_matches_reason == "":
                result.sample_form_matches_reason = error_msg
            if result.blank_form_empty_reason == "":
                result.blank_form_empty_reason = error_msg
            if result.sample_form_filled_reason == "":
                result.sample_form_filled_reason = error_msg
            if result.sample_info_masked_reason == "":
                result.sample_info_masked_reason = error_msg
